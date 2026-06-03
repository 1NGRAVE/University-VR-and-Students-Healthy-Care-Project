"""Generate the full paper as a Word document with embedded figures and tables."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============ Page Setup ============
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============ Styles ============
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_cn_paragraph(doc, text, font_name='宋体', font_size=Pt(12), bold=False,
                     alignment=None, first_line_indent=None, space_after=Pt(0),
                     space_before=Pt(0), line_spacing=1.5):
    """Add a Chinese paragraph with proper formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p

def add_heading_cn(doc, text, level=1):
    """Add a heading with Chinese font support."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_figure(doc, image_path, caption, width_inches=5.5):
    """Add a figure with caption."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        run = caption_p.add_run(caption)
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.color.rgb = RGBColor(0x63, 0x6e, 0x72)
    else:
        add_cn_paragraph(doc, f'[图片缺失: {caption}]', font_size=Pt(9))

def add_table_cn(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True
        # Header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # Alternating row colors
            if r % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # space after table
    return table

# ==================== PAPER CONTENT ====================

# ---- Title ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(4)
title_p.paragraph_format.space_before = Pt(24)
run = title_p.add_run('基于多模态生理感知与AI决策的\nVR身心脑协同训练系统设计')
run.font.size = Pt(22)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

# ---- Subtitle ----
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(12)
run = sub_p.add_run('——面向大学生群体的创新训练方案')
run.font.size = Pt(14)
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
run.font.color.rgb = RGBColor(0x63, 0x6e, 0x72)

# ---- Authors ----
author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_after = Pt(2)
run = author_p.add_run('刘  原')
run.font.size = Pt(14)
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# ---- Affiliation ----
affil_p = doc.add_paragraph()
affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil_p.paragraph_format.space_after = Pt(16)
run = affil_p.add_run('（大连理工大学 计算机科学与技术学院，辽宁 大连 116024）')
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ---- Abstract ----
add_heading_cn(doc, '摘  要', level=1)

add_cn_paragraph(doc,
    '大学生群体普遍面临久坐少动、心理压力增大与认知疲劳等多维健康问题，传统体育训练模式以单一身体锻炼为主，'
    '缺乏对心理情绪状态与认知功能的同步考量，难以实现个性化、精准化的训练指导。本文设计了一套融合虚拟现实（VR）、'
    '多模态生理传感与人工智能（AI）决策的闭环训练系统——VR身心脑协同训练系统。'
    '系统采用"传感采集—信号处理—AI分析—VR反馈"四层闭环架构，包含两套可渐进升级的硬件方案'
    '（低成本方案约600元，高精度方案约15,000元），以及一个六层递进式可解释AI分析管道。'
    '硬件层以Polar H10心率胸带为核心传感器，通过"一源多用"策略从心率变异性（HRV）信号中同时提取身体、心理与脑功能'
    '三维指标；AI层以随机森林、支持向量回归等可解释机器学习方法为核心，实现从生理状态评估到VR场景适配的全链条'
    '智能决策；VR层基于Unity与Meta Quest 3构建四类运动场景，实现心率色环、脑电柱状图、GSR压力计等实时生物反馈可视化。'
    '系统设计在理论层面具备可行性与技术创新性，为人工智能赋能大学生体育训练提供了完整的工程技术范式。',
    first_line_indent=Cm(0.74), font_size=Pt(10.5))

add_cn_paragraph(doc,
    '关键词：虚拟现实；人工智能；多模态生理感知；身心脑协同训练；生物反馈；大学生',
    font_size=Pt(10.5), bold=True, space_before=Pt(6))

# ---- English Abstract ----
add_heading_cn(doc, 'Abstract', level=1)

def add_en_paragraph(doc, text, font_size=Pt(10.5), bold=False, alignment=None,
                     first_line_indent=None, space_after=Pt(0), space_before=Pt(0)):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p

add_en_paragraph(doc,
    'College students increasingly face multidimensional health challenges including sedentary behavior, '
    'psychological stress, and cognitive fatigue. This paper presents the design of a VR Body-Mind-Brain '
    'Synergistic Training System that integrates virtual reality, multimodal physiological sensing, and '
    'artificial intelligence decision-making into a closed-loop training architecture. The system features '
    'two progressively upgradable hardware configurations, a six-tier interpretable AI analysis pipeline '
    'employing Random Forest classification, Support Vector Regression, and rule-based decision fusion, '
    'and four VR exercise scenes with real-time biofeedback visualization. The "one-source-multi-use" '
    'strategy extracts body, mind, and brain indicators from a single HRV signal, significantly lowering '
    'the hardware barrier. This design provides a complete technical paradigm for AI-empowered college '
    'student sports training innovation.',
    first_line_indent=Cm(0.74), font_size=Pt(10.5))

add_en_paragraph(doc,
    'Keywords: virtual reality; artificial intelligence; multimodal physiological sensing; '
    'body-mind-brain training; biofeedback; college students',
    font_size=Pt(10.5), bold=True, space_before=Pt(6))

# ==================== 1. 引言 ====================
doc.add_page_break()
add_heading_cn(doc, '1  引言', level=1)

add_heading_cn(doc, '1.1  大学生身心健康现状', level=2)

add_cn_paragraph(doc,
    '大学生群体的身心健康状况已成为亟待关注的公共卫生议题。中国学生体质健康调研数据显示，'
    '大学生体质健康水平连续多年呈下降趋势，不及格率在各类学生群体中居首[1]。与此同时，'
    '大学生心理健康问题日益突出——一项涵盖全国多所高校的调查显示，约20%～30%的大学生存在不同程度的'
    '焦虑、抑郁等心理困扰[2]。《中国国民心理健康发展报告（2021～2022）》指出，青年群体是'
    '心理健康风险的高发人群，18～25岁年龄段心理问题检出率显著高于其他年龄段[3]。Yang等（2023）对10项研究'
    '（样本量39,017人）的meta分析发现，体力活动与中国大学生抑郁（OR=0.69）和焦虑（OR=0.71）呈显著负相关，'
    'COVID-19期间焦虑检出率约为24.9%[30]。Shu等（2023）的meta分析（9项研究，1,613名被试）表明，有氧运动'
    '可显著提升大学生身体自尊（WMD=1.32～1.62），且90分钟/次、持续16周的干预方案效果最强[31]。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '从生活方式角度分析，大学生日均久坐时间普遍超过8小时，体力活动水平显著不足[4]。'
    '电子设备的过度使用进一步加剧了久坐行为、睡眠不足与认知疲劳的恶性循环。世界卫生组织（WHO）建议'
    '18～64岁成年人每周进行至少150～300分钟中等强度有氧运动，然而我国大学生中达到该推荐量的比例不足30%[5]。'
    '传统的高校体育课程以统一化教学为主，存在训练模式单一、个性化不足、对心理与认知维度关注缺失等局限，'
    '难以有效应对大学生群体"身（生理）—心（心理）—脑（认知）"多维健康需求。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '1.2  相关技术研究进展', level=2)

add_cn_paragraph(doc,
    '近年来，虚拟现实（VR）技术在运动训练与康复领域展现出广阔的应用前景。VR沉浸式环境能够增强运动参与感'
    '与动机，多项随机对照试验的系统综述表明，VR辅助训练在改善平衡能力、步态功能与运动依从性方面优于传统训练方式'
    '[6,7]。在VR结合生物反馈的前沿方向，Morone等（2024）的系统综述纳入11项RCT，证实VR训练对慢性下背痛的疼痛、'
    '功能障碍和运动恐惧症均有显著改善效果[32]。Kober等（2025）综述了31项VR-神经反馈研究，发现VR反馈相比2D屏幕'
    '反馈在用户体验和心理效果方面具有显著优势[33]。在生物反馈领域，心率变异性（HRV）已被确立为评估自主神经系统'
    '功能与心理压力的黄金标准——Task Force of '
    'the European Society of Cardiology and the North American Society of Pacing and Electrophysiology '
    '于1996年发布的《心率变异性：测量标准、生理学解释与临床应用》至今已被引用超过15,000次，为HRV的科研与'
    '临床应用提供了权威的方法学指导[8]。Polar H10胸带式心率传感器经多项独立验证研究确认，其静息状态下线性HRV'
    '指标（RMSSD、pNN50）与12导联ECG的一致性达到ICC=0.90～1.00水平，MAPE低至2.16%[9,10]。然而，在高强度动态'
    '运动中（如障碍赛、循环训练），Polar H10的信号丢失率可升至29.9%～40.4%，非线性HRV指标（如DFA-α1）的一致性'
    '显著下降[10a]。这提示在VR高强度训练场景中需配套信号质量实时监测与自动插值/剔除策略。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '在情绪识别领域，Russell于1980年提出的情绪环状模型（Circumplex Model of Affect）将情绪组织在效价（Valence）'
    '与唤醒度（Arousal）两个正交维度上，为基于生理信号的情绪量化提供了理论框架[11]。Davidson的前额叶EEG不对称模型'
    '揭示了前额叶Alpha波段（8～13 Hz）的左右半球不对称性与情绪效价之间的关联：左侧前额叶Alpha功率降低与'
    '趋近动机和积极情绪相关，右侧前额叶Alpha功率降低与回避动机和消极情绪相关[12,13]。需指出的是，Harmon-Jones与Allen'
    '（1998）的后续研究对"效价-不对称性"的简单映射提出了重要修正——愤怒虽属负效价但具有趋近动机，同样表现为左侧前额叶'
    '激活[12a]。近期大规模重复性研究（#EEGManyLabs）发现静息态FAA与心理病理学之间的效应量较小，提示FAA作为'
    '个体差异稳定生物标志物的可靠性需谨慎看待[12b]。因此，本系统将FAA作为情绪效价的"实时状态指示器"（而非稳定特质'
    '标志物），结合GSR皮肤电导与LF/HF进行多指标交叉验证，以提高情绪评估的稳健性。'
    '在注意力评估方面，Theta/Beta比值（TBR）作为注意力相关脑电指标已有大量研究——Lubar最早报道了注意力'
    '缺陷障碍患者的TBR升高现象[14]，Monastra等对其临床应用进行了初步验证[15]。Arns等的大规模meta分析（2013）'
    '对9项研究（1,253例ADHD vs. 517例对照）进行了系统评价，发现TBR的grand-mean效应量d=0.75（6～13岁）和'
    'd=0.62（6～18岁），但效应量异质性显著。作者明确指出TBR"不能被视为可靠的诊断指标"，更适合作为"预后性'
    '而非诊断性"工具使用，即对治疗分层具有参考价值而对诊断分类价值有限[16]。2024年更新meta分析进一步表明，'
    '在控制慢速Alpha峰值频率后，TBR效应量降至非显著水平（d从−0.212到+0.218），确认了TBR在诊断方面的局限性[16a]。'
    '因此，本系统将TBR作为注意力状态的"趋势监测"指标而非绝对诊断标准，在连续监测中通过个体内基线比较'
    '（TBR升高>1.0标准差）触发注意力下降预警。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '在可穿戴传感与信号处理工具链方面，Makowski等（2021）开发的NeuroKit2为Python生态提供了从ECG/PPG/EDA/EEG'
    '信号处理到高级生理特征提取的一站式解决方案，其ECG R峰检测基于Pan-Tompkins算法，在MIT-BIH心律失常数据库上'
    '达到超过99%的准确率[17]。Gramfort等（2013）开发的MNE-Python已成为脑电图/脑磁图分析的科研标准工具，'
    '提供了独立成分分析（ICA）去伪影、Welch功率谱密度估计等关键功能[18]。这些开源工具为本系统的信号处理管道'
    '提供了坚实的技术基础。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '在人工智能赋能体育训练方面，可解释人工智能（Explainable AI, XAI）正成为健康与运动科学领域的重要趋势。'
    '与传统黑盒深度学习模型相比，基于规则与经典机器学习的可解释决策系统在健康干预场景中具有更高的可信度、'
    '安全性与临床可接受性[19]。Kranzinger等（2025）对19项运动科学XAI研究的系统综述发现，SHAP是目前运动科学中'
    '占绝对主导地位的XAI方法，应用涵盖游泳成绩预测（R²=0.93）、运动损伤风险识别和战术决策支持等领域[34]。'
    'Abdelaal等（2024）综述了25项穿戴设备数据分析中的XAI研究，指出仅约20%包含用户评估，缺乏领域专家对解释的'
    '验证仍是主要研究空白[35]。Shen等（2024）首次全面综述了面向生理信号（ECG、EEG、EDA、PPG）的XAI方法，'
    '为多模态生理信号的可解释分析提供了方法论指引[36]。将多模态生理传感、可解释AI决策与VR沉浸式训练三者'
    '整合为一个闭环系统，是当前运动科学与人机交互领域的创新方向，相关系统设计研究仍然较少。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '1.3  研究目标', level=2)

add_cn_paragraph(doc,
    '针对上述研究空白，本文提出并设计了一套"VR身心脑协同训练系统"。系统的核心设计目标包括：（1）构建多模态'
    '生理数据采集管道，实现从"身、心、脑"三个维度同步感知用户实时状态；（2）搭建六层可解释AI分析管道，'
    '将原始生理信号转化为个性化运动训练推荐；（3）设计沉浸式VR生物反馈交互界面，实现训练过程的实时自适应闭环。'
    '本文聚焦于系统的工程技术方案设计与方法论论证，为后续的系统实现与实证研究奠定基础。',
    first_line_indent=Cm(0.74))

# ==================== 2. 系统总体架构 ====================
add_heading_cn(doc, '2  系统总体架构', level=1)

add_cn_paragraph(doc,
    'VR身心脑协同训练系统采用"四层闭环"架构，自底向上依次为：数据采集层、信号处理层、AI分析引擎层与VR渲染交互层。'
    '各层之间通过Lab Streaming Layer（LSL）实时流协议与JSON/HTTP接口进行通信，形成从传感器到VR场景的完整数据闭环。'
    '图1展示了系统的总体架构。',
    first_line_indent=Cm(0.74))

add_figure(doc, 'docs/paper_figures/fig1_system_architecture.png',
           '图1  VR身心脑协同训练系统总体架构')

add_cn_paragraph(doc,
    '系统设计遵循"渐进式复杂度"原则——提供两套可相互兼容的硬件方案：低成本方案（方案一）仅需约600元人民币，'
    '以单一Polar H10胸带为核心传感器，通过信号处理算法从HRV中挖掘多维度信息；高精度方案（方案二）增加Muse S脑电'
    '头带与Empatica E4腕带，总成本约15,000元，实现每个维度的直接生理测量。两套方案共用同一套数据管道架构与AI分析'
    '框架，支持从方案一到方案二的无缝渐进式升级。',
    first_line_indent=Cm(0.74))

# ==================== 3. 硬件方案与数据采集 ====================
add_heading_cn(doc, '3  硬件方案与多模态数据采集', level=1)

add_heading_cn(doc, '3.1  低成本方案（方案一）', level=2)

add_cn_paragraph(doc,
    '方案一的设计哲学是"一源多用"（One-Source-Multi-Use）——利用Polar H10胸带输出的高精度RR间期序列（采样率130 Hz），'
    '通过信号处理同时提取身、心、脑三个维度的生理指标。该策略的关键科学依据在于：HRV是反映自主神经系统调控功能的'
    '综合窗口——时域指标RMSSD反映副交感神经（迷走神经）活性（身→心），频域指标LF/HF比值反映交感-副交感平衡（心），'
    '非线性指标样本熵反映中枢-自主神经网络的调控灵活性与复杂度（心→脑）[8]。Polar H10采用医用级不锈钢电极与硅胶'
    '防滑胸带，在运动状态下仍能保持高质量的ECG信号采集，其RR间期精度已在多项独立验证研究中得到确认[9,10]。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '此外，方案一充分利用了VR头显（Meta Quest 3）自带的传感能力作为补充数据源：6DoF头部追踪提供头部运动平滑度与'
    '姿势稳定性指标，23关节手部追踪提供运动轨迹与关节活动范围，手柄IMU提供加速度与角速度数据。VR头显内置的认知'
    '测试模块（Stroop色词干扰任务、N-back工作记忆任务、简单反应时任务）为脑功能的"行为学代理指标"提供了重要补充。'
    '方案一共提取约12维特征，覆盖身、心、脑三个维度的核心评估需求。表1总结了方案一的数据采集与特征提取方案。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc, '表1  方案一（低成本）数据采集与特征提取方案', font_size=Pt(9), bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))

add_table_cn(doc,
    ['传感器', '原始数据', '采样率', '提取特征', '对应维度', '生理含义'],
    [
        ['Polar H10', 'ECG波形', '130 Hz', 'mean_hr, hr_zone', '身', '实时运动负荷'],
        ['Polar H10', 'RR间期序列', '逐拍', 'SDNN, RMSSD, pNN50', '心', '副交感活性/压力水平'],
        ['Polar H10', 'RR间期序列', '逐拍', 'LF power, HF power, LF/HF', '心', '交感-副交感平衡'],
        ['Polar H10', 'RR间期序列', '逐拍', 'Sample Entropy', '脑(代理)', '自主神经调控复杂度'],
        ['Polar H10', 'RR间期序列', '逐拍', 'Baevsky SI', '心', '综合压力指数'],
        ['VR 6DoF追踪', '头部位置+旋转', '~90 Hz', 'head_smoothness, posture_stability', '身', '运动控制质量'],
        ['VR 手部追踪', '23关节3D坐标', '~30 Hz', 'trajectory_length, joint_angle_range', '身', '动作幅度/完成度'],
        ['VR 手柄IMU', '加速度+角速度', '~100 Hz', 'acc_peak, movement_freq', '身', '运动爆发力/节奏'],
        ['VR 认知测试', '行为响应', '按需', 'reaction_time, stroop_interference', '脑(代理)', '执行功能/处理速度'],
        ['VR 认知测试', '行为响应', '按需', 'nback_d_prime, attention_decay', '脑(代理)', '工作记忆/注意力'],
    ])

add_heading_cn(doc, '3.2  高精度方案（方案二）', level=2)

add_cn_paragraph(doc,
    '方案二在方案一的基础上增加了两个专用传感器，实现各维度的直接生理测量：（1）Muse S脑电头带——4通道干电极EEG'
    '（AF7、AF8、TP9、TP10，采样率256 Hz），提供前额叶与颞顶叶皮层电活动数据，支持基于Davidson前额叶不对称模型'
    '的情绪效价量化[12]、基于Theta/Beta比值的注意力量化[14-16]，以及Alpha峰值频率（IAF）的认知处理速度评估[20]；'
    '（2）Empatica E4腕带——提供皮肤电导（GSR/EDA，4 Hz）、血容量脉搏（PPG，64 Hz）、皮肤温度（4 Hz）与3轴加速度'
    '（32 Hz），支持基于cvxEDA算法的皮肤电导tonic/phasic分解[21]，以及基于Russell模型的高精度情绪唤醒度评估[11]。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc, '表2  两套硬件方案的对比', font_size=Pt(9), bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))

add_table_cn(doc,
    ['对比维度', '方案一（低成本）', '方案二（高精度）'],
    [
        ['传感器', 'Polar H10 + VR内置', 'Polar H10 + Muse S + Empatica E4 + VR内置'],
        ['总成本', '~¥600 ($80)', '~¥15,000 ($2,130)'],
        ['特征维度', '~12维', '~31维'],
        ['身体评估', 'HR + HRV（直接）★★★★★', 'HR + HRV（直接）★★★★★'],
        ['心理评估', 'HRV间接指标 ★★★★', 'EEG + GSR直接测量 ★★★★★'],
        ['脑功能评估', 'VR认知测试（行为代理）★★★', 'EEG直接 + 认知测试 ★★★★★'],
        ['情绪粒度', '压力/放松 二元判断', '效价-唤醒度 连续坐标定位'],
        ['注意力评估', '间接推断', 'Theta/Beta比直接量化'],
        ['适用场景', '日常训练、大规模推广', '科研级精细评估、临床康复'],
    ])

# ==================== 4. AI分析引擎 ====================
add_heading_cn(doc, '4  AI分析引擎设计', level=1)

add_cn_paragraph(doc,
    'AI分析引擎是系统的核心智能层，采用六层递进式可解释管道（Six-Tier Progressive Interpretable Pipeline）架构。'
    '与端到端的深度学习黑盒模型不同，本设计选择"规则引擎+经典机器学习"的混合策略，确保每一层的输入、处理逻辑与输出'
    '完全可追溯、可解释——这在涉及人体生理安全的健康干预场景中至关重要[19]。图2展示了六层管道的完整架构与数据流。',
    first_line_indent=Cm(0.74))

add_figure(doc, 'docs/paper_figures/fig2_ai_pipeline.png',
           '图2  六层递进式AI分析管道架构')

add_heading_cn(doc, '4.1  Layer 1——生理状态评估', level=2)

add_cn_paragraph(doc,
    'Layer 1接收经过NeuroKit2[17]和HeartPy预处理的ECG/HRV特征向量、VR运动追踪数据以及方案二中Empatica E4的PPG/'
    '温度/加速度数据，输出四项0—100分的子评分：（1）心脏负荷评分（cardiac_load）——基于当前心率占最大心率百分比与'
    'HR储备利用率；（2）自主神经平衡评分（autonomic_balance）——综合RMSSD、LF/HF比值与Baevsky压力指数，RMSSD < 20 ms'
    '提示高压力/恢复不足，LF/HF > 3提示交感神经过度激活；（3）循环代谢评分（circulatory）——基于PPG波形形态与'
    '皮肤温度趋势；（4）运动质量评分（movement_quality）——基于VR追踪数据的运动平滑度、左右对称性与目标动作匹配度。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '4.2  Layer 2——心理情绪评估', level=2)

add_cn_paragraph(doc,
    'Layer 2基于Russell的情绪环状模型[11]，将用户的实时情绪状态映射到以效价（Valence, X轴，-1～+1）和唤醒度'
    '（Arousal, Y轴，0～1）为坐标的二维情绪空间。效价轴以EEG额叶Alpha不对称性为核心指标——通过计算不对称分数'
    'α_asym = (AF8_alpha − AF7_alpha)/(AF8_alpha + AF7_alpha)，正值为负效价（右前额活跃→回避），负值为正效价'
    '（左前额活跃→趋近）[12,13]。唤醒度轴以GSR皮肤电导tonic水平为核心指标，辅以GSR phasic峰值频率与LF/HF比值'
    '进行交叉验证。对于方案一（无EEG与GSR），Layer 2降级为基于HRV RMSSD与LF/HF的简化压力-放松二元判断，'
    '仍能提供基本的情绪状态参考。图3展示了情绪环状模型与生理信号映射关系。',
    first_line_indent=Cm(0.74))

add_figure(doc, 'docs/paper_figures/fig3_emotion_model.png',
           '图3  Russell情绪环状模型与多模态生理信号映射')

add_heading_cn(doc, '4.3  Layer 3——认知状态评估', level=2)

add_cn_paragraph(doc,
    'Layer 3通过VR内置认知测试模块评估用户的实时认知功能。Stroop色词干扰任务测量执行功能与抑制控制能力——干扰效应'
    '（incongruent RT − congruent RT）< 50 ms为执行功能优秀，50～120 ms为正常，> 120 ms提示前额叶执行资源显著下降[22]。'
    'N-back工作记忆任务（2-back条件）通过信号检测论d\'敏感度指标量化工作记忆容量——d\' > 2为优秀，1～2为正常，'
    '< 1提示工作记忆负荷已满或认知疲劳[23]。简单反应时任务（RT）提供认知处理速度的基线指标——RT < 250 ms为快速，'
    '250～350 ms为正常，> 350 ms提示中枢疲劳[24]。方案二中，Layer 3还融合EEG Gamma波段（30～45 Hz）功率作为执行'
    '加工活跃度的神经指标[25]，以及Theta/Beta比值作为注意力的跨任务验证指标[16]。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '4.4  Layer 4——综合状态融合', level=2)

add_cn_paragraph(doc,
    'Layer 4将前三个Layer的独立评估结果融合为一个全局状态向量（Global State Vector）。融合采用可解释的规则化方法'
    '而非黑盒神经网络——这意味着每个融合权重的物理含义都是明确的。除各维度的评分外，该层还计算三个关键的跨域耦合'
    '指标：（1）身心耦合度——当自主神经平衡评分与放松深度评分偏差 < 15分时，判定为"身心同步"状态；偏差 > 15分且'
    '方向不一致时触发针对性干预提示；（2）脑身耦合度——当认知功能评分低而身体负荷评分高时，可能提示中枢疲劳主导的'
    '运动能力下降；（3）恢复准备度——综合RMSSD趋势、GSR tonic趋势、皮肤温度趋势与认知功能趋势的加权评分，用于判断'
    '用户是否适合进行高强度训练[8,26]。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '4.5  Layer 5——推荐决策引擎', level=2)

add_cn_paragraph(doc,
    'Layer 5是AI的核心决策层，采用四级级联决策流程。第一级——安全规则门控（不可被ML覆盖）：当HR > 90% HRmax时触发'
    '"立即停止运动"协议，当RMSSD < 基线×0.5时限制为仅允许恢复类活动，当皮肤温度上升斜率 > 1.0 °C/min持续3分钟时'
    '触发热应激预警。第二级——生理边界过滤：根据自主神经平衡、心脏负荷、循环代谢与运动质量评分，通过确定性规则筛选'
    '候选运动类型集合。第三级——机器学习分类与回归：使用scikit-learn实现的随机森林（Random Forest, RF）分类器进行'
    '运动类型概率预测[27]，使用支持向量回归（Support Vector Regression, SVR, RBF核）进行基础运动强度（0—100%）预测。'
    '第四级——情绪/认知上下文修正：根据Layer 2的情绪象限与Layer 3的注意力/认知评分，对ML输出进行乘法因子微调'
    '（每个因子0.7～1.15），例如负效价情绪时强度降至85%，注意力不足（< 30分）时强度降至80%，高唤醒正效价'
    '（兴奋状态）时强度可适度提升至115%。最终的强度值经clamp(5, 95)确保始终在安全范围内。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '4.6  Layer 6——VR适配输出', level=2)

add_cn_paragraph(doc,
    'Layer 6将Layer 5的抽象推荐转化为具体的VR场景参数，输出包含：运动场景选择（海滩/山湖/森林/抽象空间）、'
    '背景音频方案（432 Hz / 528 Hz / 120–160 BPM节奏）、视觉引导强度、环境自适应参数（天空色温范围、粒子密度、'
    '光照强度）以及实时触发条件（如"连续2个epoch期间RMSSD下降 > 20%时VR教练语音提醒放慢节奏"）[28]。'
    '最终输出为JSON格式的AI推荐包，通过HTTP :5000端口发送至Unity VR前端，同时通过LSL流推送至数据记录器。',
    first_line_indent=Cm(0.74))

# ==================== 5. VR交互设计 ====================
add_heading_cn(doc, '5  VR交互与生物反馈设计', level=1)

add_heading_cn(doc, '5.1  运动场景设计', level=2)

add_cn_paragraph(doc,
    '系统基于Unity 2022 LTS与Meta XR SDK/OpenXR构建四类VR运动场景，分别对应不同的训练目标与情绪状态适配。'
    '图4展示了四类场景的概览与生物反馈元素映射。',
    first_line_indent=Cm(0.74))

add_figure(doc, 'docs/paper_figures/fig4_vr_scenes.png',
           '图4  四类VR运动场景与生物反馈元素')

add_cn_paragraph(doc, '表3  四类VR运动场景的设计参数与适配策略', font_size=Pt(9), bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))

add_table_cn(doc,
    ['场景', '训练类型', '目标情绪象限', '背景音频', '核心视觉引导', 'HUD元素', '环境自适应'],
    [
        ['海滩日落', '呼吸训练\n(4-4-4-4方呼吸)', '高唤醒-负效价\n(紧张/焦虑)', '432 Hz\n+ 海浪声', '呼吸光环缩放\n吸气扩大/呼气缩小', 'HR色环\nGSR水滴\n呼吸计时', '天空色温暖橙↔深蓝\n压力↓→暖色过渡'],
        ['山湖', '太极引导\n(低强度/精细动作)', '低唤醒-正效价\n(平静/满足)', '528 Hz\n+ 自然环境声', '虚拟化身演示\n目标球引导手部', '关节活动指示\n运动对称性环\n动作完成度%', '暖黄稳定光照\n注意力↑→亮度↑'],
        ['森林小径', '正念漫步\n(中等强度/有氧)', '低唤醒-负效价\n(疲劳/低落)', '528 Hz + 鸟鸣\n渐进120 BPM', 'Alpha功率光晕\n步频节奏提示', 'HR区间环\nEEG Alpha光晕\n步频指示器', '翠绿暖光\nAlpha↑→萤火虫粒子↑'],
        ['抽象空间', '高强度挑战\n(间歇/HIIT)', '高唤醒-正效价\n(兴奋/激动)', '140-160 BPM\n电子节奏', '动态目标球\n连击计分系统', 'HR区间环(大字)\n连击计数器\n剩余时间', '霓虹色快速变幻\nHR↑→节奏加速'],
    ])

add_heading_cn(doc, '5.2  实时生物反馈可视化', level=2)

add_cn_paragraph(doc,
    '系统在VR视野边缘渲染五个核心生物反馈HUD元素：（1）HR区间色环——以脚部/腕部彩色光环显示实时心率区间，'
    '蓝色（Zone 1, <60% HRmax）→绿色（Zone 2, 60-70%）→黄色（Zone 3, 70-80%）→橙色（Zone 4, 80-90%）→'
    '红色（Zone 5, >90%），光环的脉动频率与实时心跳同步[29]。（2）EEG频段柱状图——以Delta/Theta/Alpha/Beta/Gamma'
    '五个频段的实时功率柱状图显示脑电活动，柱体颜色从冷色调（低频）过渡到暖色调（高频）。（3）GSR压力水滴——'
    '以填充/排空的水滴动画显示皮肤电导水平，满水滴代表高交感激活/压力，排空状态代表放松。（4）呼吸引导光球——'
    '以4秒扩大-4秒保持-4秒缩小-4秒保持的动画光球引导腹式呼吸节奏。（5）AI教练文字提示——在VR视野上方以半透明'
    '文字显示当前推荐的运动类型、强度与个性化建议（如"放慢节奏，注意呼吸"）。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '5.3  环境自适应机制', level=2)

add_cn_paragraph(doc,
    'VR环境的视觉与音频参数根据用户实时生理状态进行连续自适应调节，形成超越显式HUD的隐性生物反馈通道：'
    '（1）天空色温——根据Baevsky压力指数在暖橙（放松，<50）与冷灰（高压，>500）之间连续过渡；（2）背景音频频率——'
    '根据HR区间在432 Hz（静息/恢复）与528 Hz（活跃/专注）之间调整，高强度场景下切换为BPM与实时心率同步的节奏音频；'
    '（3）粒子密度——根据EEG Alpha功率动态增减萤火虫/花瓣等放松感粒子，Alpha功率越高粒子越密集；'
    '（4）光照强度——根据注意力评分（TBR）动态调节，注意力越高场景亮度越大，注意力涣散时光线柔和变暗以减少认知负荷。'
    '表4总结了环境自适应参数的调节逻辑。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc, '表4  环境自适应参数调节逻辑', font_size=Pt(9), bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))

add_table_cn(doc,
    ['环境参数', '触发生理指标', '调节逻辑', '自适应效果'],
    [
        ['天空色温', 'Baevsky压力指数', 'SI < 50 → 暖橙; SI 50-150 → 中性; SI > 500 → 冷灰', '压力视觉氛围调节'],
        ['背景音频频率', 'HR区间 + 情绪象限', 'Zone1-2 → 432Hz; Zone3-4 → 528Hz; Zone5 → BPM同步', '音频-生理节律同步'],
        ['粒子密度', 'EEG Alpha功率', 'Alpha > 200%基线 → 最大密度; < 100%基线 → 无粒子', '放松状态视觉增强'],
        ['光照强度', 'Theta/Beta比值', 'TBR < 1.5 → 高亮度; TBR > 2.5 → 柔和降低', '注意力状态适配'],
        ['VR教练语音', 'RMSSD趋势 + GSR峰值', 'RMSSD↓>20% → 提醒放慢; GSR↑突增 → 安抚提示', '实时安全与情绪干预'],
    ])

# ==================== 6. 讨论 ====================
add_heading_cn(doc, '6  讨论', level=1)

add_heading_cn(doc, '6.1  系统创新点分析', level=2)

add_cn_paragraph(doc,
    '本系统的设计在以下四个方面体现了技术创新性：',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '第一，"一源多用"的多维感知策略。传统多模态生理感知系统通常要求每个评估维度配置独立传感器，导致硬件成本高昂、'
    '系统复杂度大、佩戴舒适度低。本系统提出的"一源多用"策略从HRV信号中同时提取身体（心率负荷）、心理（RMSSD压力、'
    'LF/HF情绪唤醒）与脑功能（样本熵神经调控复杂度）三个维度的指标，将单一传感器的信息利用效率最大化。这一策略的'
    '科学基础在于HRV是自主神经系统功能的多维窗口——时域、频域与非线性指标各自携带不同维度的生理信息[8]。在低成本'
    '方案（约600元）下，该策略使系统具备了准入门槛极低的身心脑三维评估能力，为大规模校园推广应用创造了条件。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '第二，六层可解释AI管道。目前运动训练领域的AI应用多为端到端深度学习模型，虽然预测精度较高但决策过程不透明，'
    '存在安全性隐患[19]。本系统采用的"安全规则门控→生理边界过滤→RF分类→SVR回归→上下文微调"级联架构，在每一层'
    '都保持了完全的可解释性——用户可以追踪任何一条推荐决策的完整推理路径。安全规则层独立于ML模型，确保了生理安全'
    '约束的绝对优先级。这种设计特别适合健康干预场景，因为当AI推荐与临床判断冲突时，可解释性使人工审核与干预成为可能。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '第三，闭环VR生物反馈机制。传统VR运动训练多为"预设内容+被动体验"模式，环境参数固定不变。本系统通过LSL实时'
    '流与HTTP API实现了从传感器到VR场景的完整数据闭环，使VR环境成为"活的"自适应训练空间——天空色温、音频频率、'
    '粒子密度、光照强度等环境参数均根据实时生理状态连续调节。这种隐性生物反馈（ambient biofeedback）与显性HUD'
    '元素（心率环、EEG柱状图等）形成双层反馈通道，使用户在自然沉浸中获得个性化的训练引导。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '第四，渐进式硬件架构。方案一与方案二的共享数据管道设计保证了系统可以从低成本方案平滑升级到高精度方案，'
    '无需重新开发数据采集与AI分析模块。这种"最低可行产品（MVP）优先+渐进增强"的工程策略降低了系统开发风险，'
    '也为不同预算与应用场景（课堂教学、科研实验、临床康复）提供了灵活的部署选项。',
    first_line_indent=Cm(0.74))

add_heading_cn(doc, '6.2  与现有研究的比较', level=2)

add_cn_paragraph(doc, '表5  本系统与现有VR运动训练系统的比较', font_size=Pt(9), bold=True,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))

add_table_cn(doc,
    ['特征', '传统VR运动游戏\n(e.g., Beat Saber)', 'VR运动康复系统\n(现有研究)', 'VR生物反馈训练\n(现有研究)', '本系统'],
    [
        ['生理传感', '仅手柄IMU', 'HR监测（可选）', '1-2种传感器', '多模态（HRV/EEG/GSR/PPG）'],
        ['心理评估', '无', '无', '有限（单一指标）', '效价-唤醒度二维情绪定位'],
        ['认知评估', '无', '无', '无', 'Stroop/N-back/反应时'],
        ['AI个性化', '固定难度曲线', '基于HR的区间调整', '基于阈值的简单规则', 'RF+SVR+多层规则级联'],
        ['AI可解释性', 'N/A', 'N/A', '高（简单规则）', '高（全管道可追溯）'],
        ['VR环境自适应', '无', '有限', '1-2个参数', '5+环境参数连续自适应'],
        ['成本', '~$500（仅VR）', '不等', '不等', '$80起（不含VR头显）'],
        ['目标人群', '普通消费者', '康复患者', '特定训练人群', '大学生群体（普适+个性化）'],
    ])

add_heading_cn(doc, '6.3  局限性与未来工作', level=2)

add_cn_paragraph(doc,
    '本系统目前处于详细设计阶段，存在以下主要局限：（1）尚未完成系统编码实现与硬件集成测试，各层模块的技术可行性'
    '基于文献论证和Web模拟原型验证，仍需在实际硬件环境中进行端到端测试；（2）AI模型的训练需要基准数据集——当前设计'
    '中采用基于文献与专家知识的规则权重方案，在系统实现后需要通过5—10人×6个session的数据收集协议建立初始训练集[30]；'
    '（3）低成本方案中"脑"维度的HRV代理指标（样本熵、总功率）只能反映自主神经层面的调控复杂度，无法直接测量皮层'
    '电活动——这在高精度方案中可以通过Muse S EEG得到解决，但在低成本方案中依然是固有局限；（4）VR认知测试'
    '（Stroop/N-back）的行为学指标虽然具有成熟的心理学测量学基础，但其作为"脑功能代理"的有效性仍需通过'
    'EEG-行为联合验证研究进一步确认。',
    first_line_indent=Cm(0.74))

add_cn_paragraph(doc,
    '未来工作将沿着四个方向推进：（1）系统实现——按照六阶段实施计划（环境搭建→数据采集管道→信号处理管道→AI管道'
    '开发→VR应用开发→集成测试与验证）完成全系统编码与集成；（2）实证验证——通过随机对照试验（RCT）设计，'
    '评估系统对大学生体质健康指标、心理健康量表评分、认知功能测试成绩与运动依从性的实际干预效果；（3）算法优化——'
    '在积累足够训练数据后，探索XGBoost等梯度提升方法在运动类型分类中的性能提升潜力[31]，以及在线学习机制'
    '对模型个性化的进一步优化；（4）场景扩展——面向特定应用场景（如考试焦虑缓解、专注学习增强、睡眠改善）'
    '开发专门的VR训练模块，扩大系统的应用覆盖范围。',
    first_line_indent=Cm(0.74))

# ==================== 7. 结论 ====================
add_heading_cn(doc, '7  结论', level=1)

add_cn_paragraph(doc,
    '本文针对大学生群体"身—心—脑"多维健康需求与传统体育训练模式单一化之间的矛盾，设计了一套基于多模态生理感知、'
    '可解释AI决策与VR生物反馈的闭环协同训练系统。系统的主要贡献包括：（1）提出了"一源多用"的HRV多维信息提取策略，'
    '以极低的硬件成本（约600元）实现了身、心、脑三个维度的同步评估；（2）构建了六层递进式可解释AI分析管道，'
    '以"安全规则→生理边界→ML分类→回归预测→上下文修正→VR适配"的级联架构替代黑盒模型，确保决策透明与生理安全；'
    '（3）设计了四类VR运动场景与双层生物反馈（显性HUD + 隐性环境自适应）机制，实现了训练过程的实时个性化闭环。'
    '系统设计在理论层面具备可行性与创新性，为人工智能赋能大学生体育训练提供了完整的工程技术范式。'
    '后续研究将重点推进系统实现与随机对照试验验证。',
    first_line_indent=Cm(0.74))

# ==================== 参考文献 ====================
doc.add_page_break()
add_heading_cn(doc, '参考文献', level=1)

references = [
    '[1] 教育部体育卫生与艺术教育司. 全国学生体质健康调研结果公报[R]. 北京: 中华人民共和国教育部, 2021.',
    '[2] 傅小兰, 张侃, 陈雪峰, 等. 中国国民心理健康发展报告（2021～2022）[M]. 北京: 社会科学文献出版社, 2023.',
    '[3] 中国科学院心理研究所. 中国国民心理健康发展报告（2021～2022）[R]. 北京, 2023.',
    '[4] Tremblay MS, Aubert S, Barnes JD, et al. Sedentary Behavior Research Network (SBRN) — '
    'Terminology Consensus Project process and outcome[J]. International Journal of Behavioral '
    'Nutrition and Physical Activity, 2017, 14(1): 75.',
    '[5] World Health Organization. WHO guidelines on physical activity and sedentary behaviour[M]. '
    'Geneva: World Health Organization, 2020.',
    '[6] Rutkowski S, Kiper P, Cacciante L, et al. Use of virtual reality-based training in different '
    'fields of rehabilitation: A systematic review and meta-analysis[J]. Journal of Rehabilitation '
    'Medicine, 2020, 52(11): jrm00121.',
    '[7] Qian J, McDonough DJ, Gao Z. The effectiveness of virtual reality exercise on individual\'s '
    'physiological, psychological and rehabilitative outcomes: A systematic review[J]. International '
    'Journal of Environmental Research and Public Health, 2020, 17(11): 4133.',
    '[8] Task Force of the European Society of Cardiology and the North American Society of Pacing '
    'and Electrophysiology. Heart rate variability: standards of measurement, physiological '
    'interpretation, and clinical use[J]. Circulation, 1996, 93(5): 1043-1065. (引用次数: >15,000)',
    '[9] Gilgen-Ammann R, Schweizer T, Wyss T. RR interval signal quality of a heart rate monitor '
    'and an ECG Holter at rest and during exercise[J]. European Journal of Applied Physiology, 2019, '
    '119(7): 1525-1532.',
    '[10] Schaffarczyk M, Rogers B, Reer R, et al. Validity of the Polar H10 sensor for heart rate '
    'variability analysis during resting state and incremental exercise in recreational athletes[J]. '
    'Sensors, 2022, 22(17): 6536. DOI: 10.3390/s22176536\n'
    '[10a] Machado R, Pereira T, Silva H, et al. Signal quality and reliability of wearable devices '
    'during dynamic physical activities: A systematic review[J]. Sensors, 2025, 25(19): 6049.',
    '[11] Russell JA. A circumplex model of affect[J]. Journal of Personality and Social Psychology, '
    '1980, 39(6): 1161-1178. DOI: 10.1037/h0077714',
    '[12] Davidson RJ, Ekman P, Saron CD, et al. Approach-withdrawal and cerebral asymmetry: '
    'Emotional expression and brain physiology I[J]. Journal of Personality and Social Psychology, '
    '1990, 58(2): 330-341. DOI: 10.1037/0022-3514.58.2.330\n'
    '[12a] Harmon-Jones E, Allen JJB. Anger and frontal brain activity: EEG asymmetry consistent '
    'with approach motivation despite negative affective valence[J]. Journal of Personality and '
    'Social Psychology, 1998, 74(5): 1310-1316.\n'
    '[12b] Kaur S, Lopez-Duran N, et al. (#EEGManyLabs). Resting frontal EEG asymmetry and '
    'psychopathology: A multi-site, multi-lab registered replication[J]. Biological Psychiatry: '
    'Cognitive Neuroscience and Neuroimaging, 2025. DOI: 10.1016/j.bpsc.2024.12.005',
    '[13] Davidson RJ. What does the prefrontal cortex "do" in affect: perspectives on frontal EEG '
    'asymmetry research[J]. Biological Psychology, 2004, 67(1-2): 219-234. DOI: 10.1016/j.biopsycho.2004.03.008',
    '[14] Lubar JF. Discourse on the development of EEG diagnostics and biofeedback for '
    'attention-deficit/hyperactivity disorders[J]. Biofeedback and Self-Regulation, 1991, 16(3): 201-225.',
    '[15] Monastra VJ, Lubar JF, Linden M, et al. Assessing attention deficit hyperactivity disorder '
    'via quantitative electroencephalography: an initial validation study[J]. Neuropsychology, 1999, '
    '13(3): 424-433.',
    '[16] Arns M, Conners CK, Kraemer HC. A decade of EEG theta/beta ratio research in attention '
    'deficit/hyperactivity disorder: a meta-analysis[J]. Journal of Attention Disorders, 2013, 17(5): '
    '374-383. DOI: 10.1177/1087054712460087\n'
    '[16a] Boxum E, Voetterl H, Arns M, et al. Theta/beta ratio in ADHD: An updated meta-analysis '
    'controlling for slow alpha peak frequency[J]. Applied Psychophysiology and Biofeedback, 2024. '
    'DOI: 10.1007/s10484-024-09649-y',
    '[17] Makowski D, Pham T, Lau ZJ, et al. NeuroKit2: A Python toolbox for neurophysiological signal '
    'processing[J]. Behavior Research Methods, 2021, 53(4): 1689-1696.',
    '[18] Gramfort A, Luessi M, Larson E, et al. MEG and EEG data analysis with MNE-Python[J]. '
    'Frontiers in Neuroscience, 2013, 7: 267.',
    '[19] Arrieta AB, Díaz-Rodríguez N, Del Ser J, et al. Explainable Artificial Intelligence (XAI): '
    'Concepts, taxonomies, opportunities and challenges toward responsible AI[J]. Information Fusion, '
    '2020, 58: 82-115.',
    '[20] Klimesch W. EEG alpha and theta oscillations reflect cognitive and memory performance: '
    'a review and analysis[J]. Brain Research Reviews, 1999, 29(2-3): 169-195.',
    '[21] Greco A, Valenza G, Lanata A, et al. cvxEDA: A convex optimization approach to '
    'electrodermal activity processing[J]. IEEE Transactions on Biomedical Engineering, 2016, 63(4): '
    '797-804.',
    '[22] MacLeod CM. Half a century of research on the Stroop effect: an integrative review[J]. '
    'Psychological Bulletin, 1991, 109(2): 163-203.',
    '[23] Owen AM, McMillan KM, Laird AR, et al. N-back working memory paradigm: A meta-analysis '
    'of normative functional neuroimaging studies[J]. Human Brain Mapping, 2005, 25(1): 46-59.',
    '[24] Warm JS, Parasuraman R, Matthews G. Vigilance requires hard mental work and is stressful[J]. '
    'Human Factors, 2008, 50(3): 433-441.',
    '[25] Tallon-Baudry C, Bertrand O. Oscillatory gamma activity in humans and its role in object '
    'representation[J]. Trends in Cognitive Sciences, 1999, 3(4): 151-162.',
    '[26] Pincus SM. Approximate entropy as a measure of system complexity[J]. Proceedings of the '
    'National Academy of Sciences, 1991, 88(6): 2297-2301.',
    '[27] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python[J]. '
    'Journal of Machine Learning Research, 2011, 12: 2825-2830.',
    '[28] Slater M, Sanchez-Vives MV. Enhancing our lives with immersive virtual reality[J]. Frontiers '
    'in Robotics and AI, 2016, 3: 74.',
    '[29] Tanaka H, Monahan KD, Seals DR. Age-predicted maximal heart rate revisited[J]. Journal of '
    'the American College of Cardiology, 2001, 37(1): 153-156.',
    '[30] Yang CM, Li M, Zhang Y, et al. Physical activity and mental health among Chinese college '
    'students: A systematic review and meta-analysis[J]. Medicine, 2023, 102(49): e36524. '
    'DOI: 10.1097/MD.0000000000036524',
    '[31] Shu J, Wang Y, Liu T, et al. Effects of aerobic exercise on body self-esteem among Chinese '
    'college students: A meta-analysis[J]. PLoS ONE, 2023, 18(9): e0291045. '
    'DOI: 10.1371/journal.pone.0291045',
    '[32] Morone G, Paolucci S, Iosa M, et al. Virtual reality for the treatment of chronic low back '
    'pain: A systematic review[J]. Sensors, 2024, 24(19): 6269. DOI: 10.3390/s24196269',
    '[33] Kober SE, Wood G, Hofer D, et al. Virtual reality and neurofeedback: A systematic review '
    'of applications, effects, and challenges[J]. Applied Psychophysiology and Biofeedback, 2025. '
    'DOI: 10.1007/s10484-024-09677-8',
    '[34] Kranzinger S, Stöckl M, Lames M. Explainable artificial intelligence in sports science: '
    'A systematic scoping review[J]. Discover Artificial Intelligence, 2025. '
    'DOI: 10.1007/s44163-025-00709-8',
    '[35] Abdelaal M, Galzarano S, Chessa S, et al. Explainable AI in wearable health data analytics: '
    'A systematic review[J]. JMIR Medical Informatics, 2024, 12: e53863. DOI: 10.2196/53863',
    '[36] Shen Y, Wang J, Li X, et al. Explainable artificial intelligence for physiological signal '
    'analysis: A comprehensive review[J]. Neurocomputing, 2024, 610: 128587.',
    '[37] Chen T, Guestrin C. XGBoost: A scalable tree boosting system[C]. Proceedings of the 22nd '
    'ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016: 785-794.',
    '[38] Johansson M, Andersson E, Nilsson J, et al. Validation of the Polar H10 heart rate sensor '
    'for heart rate variability analysis during exercise[J]. Frontiers in Physiology, 2026, '
    '16: 1707318. DOI: 10.3389/fphys.2025.1707318',
]

for ref in references:
    add_cn_paragraph(doc, ref, font_size=Pt(9), line_spacing=1.3, space_after=Pt(2))

# ==================== Save ====================
output_path = 'docs/VR身心脑协同训练系统-学术论文.docx'
doc.save(output_path)
print(f'Paper saved to: {output_path}')
print(f'Total references: {len(references)}')
